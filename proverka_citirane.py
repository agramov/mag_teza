# -*- coding: utf-8 -*-
"""Одит на цитиранията: за ВСЯКО изречение с [n] проверява дали има ред в
плана с дословна извадка от източника. Изходът е списък с дупките.
Употреба: python3 proverka_citirane.py [docx] [xlsx]"""
import re, sys, unicodedata
from docx import Document
import openpyxl

DOCX = sys.argv[1] if len(sys.argv)>1 else 'MT_Aleksandar_tekst.docx'
XLSX = sys.argv[2] if len(sys.argv)>2 else 'План_проверка_източници_FINAL.xlsx'

def norm(s):
    s = unicodedata.normalize('NFKC', str(s or ''))
    s = s.replace('„','"').replace('“','"').replace('”','"').replace('’',"'")
    s = re.sub(r'[–—−]', '-', s)
    return re.sub(r'\s+', ' ', s).strip()

def words(s, n=8):
    w = re.findall(r'\w+', norm(s).lower())
    return w[:n]

# --- 1. изреченията с цитат ---
d = Document(DOCX)
start = next(k for k,p in enumerate(d.paragraphs) if p.text.strip().startswith('[1] '))
claims = []           # (абз, изречение, [номера])
for i,p in enumerate(d.paragraphs[:start]):
    t = norm(p.text)
    if not t or re.match(r'^\[\d+\]', t): continue
    for sent in re.split(r'(?<=[.!?])\s+(?=[А-ЯA-Z„"])', t):
        nums = re.findall(r'\[(\d+)(?:,[^\]]*)?\]', sent)
        if nums: claims.append((i, sent, sorted(set(int(x) for x in nums))))

# --- 2. редовете с извадки ---
wb = openpyxl.load_workbook(XLSX, data_only=True)
rows = []             # (лист, ред, номера, твърдение, извадка, статус)
SHEETS = {'Ниво 2 Твърдения': dict(num=4, cl=6, ex=13, st=10),
          'Гл.2–3 Ниво 2':    dict(num=4, cl=6, ex=10, st=7)}
for nm,c in SHEETS.items():
    if nm not in wb.sheetnames: continue
    ws = wb[nm]
    for r in range(2, ws.max_row+1):
        nums = set(int(x) for x in re.findall(r'\d+', str(ws.cell(r,c['num']).value or '')))
        rows.append((nm, r, nums, norm(ws.cell(r,c['cl']).value),
                     norm(ws.cell(r,c['ex']).value), norm(ws.cell(r,c['st']).value)))

MIN_IZRECHENIYA = 2      # изискване: извадката е поне две изречения

def quoted(ex):
    """Дословна извадка = поне MIN_IZRECHENIYA изречения в кавички.
    Отделните откъси се разделят с | и всеки се разбива на изречения;
    броят се само изречения с поне 6 думи."""
    n = 0
    for q in re.findall(r'"([^"]{20,})"', ex):
        for sent in re.split(r'(?<=[.!?;:])\s+', q):
            if len(re.findall(r'\w+', sent)) >= 6: n += 1
    return n >= MIN_IZRECHENIYA

# --- 3. съпоставяне по началните думи на твърдението ---
index = {}
for rec in rows:
    k = tuple(words(rec[3]))
    if k and len(k) >= 5: index.setdefault(k, []).append(rec)

gaps = {'БЕЗ РЕД':[], 'ИЗВАДКА ПОД 2 ИЗРЕЧЕНИЯ':[], 'НЕПОТВЪРДЕН СТАТУС':[]}
ok = 0
for ab, sent, nums in claims:
    k = tuple(words(sent))
    hit = index.get(k)
    if not hit:
        # втори опит: съвпадение по източник + припокриване на думи
        best, bs = None, 0
        sw = set(re.findall(r'\w+', sent.lower()))
        for rec in rows:
            if not (rec[2] & set(nums)): continue
            s = len(sw & set(re.findall(r'\w+', rec[3].lower())))
            if s > bs: bs, best = s, rec
        hit = [best] if best and bs >= 12 else None
    if not hit:
        gaps['БЕЗ РЕД'].append((ab, nums, sent[:150])); continue
    rec = hit[0]
    if not quoted(rec[4]):
        gaps['ИЗВАДКА ПОД 2 ИЗРЕЧЕНИЯ'].append((ab, nums, rec[0], rec[1], sent[:110])); continue
    if not rec[5].startswith('Потвърдено'):
        gaps['НЕПОТВЪРДЕН СТАТУС'].append((ab, nums, rec[0], rec[1], rec[5], sent[:90])); continue
    ok += 1

print('ФАЙЛ:', DOCX)
print('изречения с цитат: %d | покрити с дословна извадка: %d | дупки: %d'
      % (len(claims), ok, len(claims)-ok))
for k,v in gaps.items():
    print('\n### %s — %d' % (k, len(v)))
    for g in v: print('  ', g)
